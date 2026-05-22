# analyzer/services/parser.py - FIXED VERSION
import PyPDF2
import docx
import os
import tempfile
from django.core.files.uploadedfile import UploadedFile, InMemoryUploadedFile

class ResumeParser:
    """Parse PDF and DOCX resume files"""
    
    @staticmethod
    def parse_resume(file):
        """Parse uploaded resume file"""
        if file.name.endswith('.pdf'):
            return ResumeParser.parse_pdf(file)
        elif file.name.endswith('.docx'):
            return ResumeParser.parse_docx(file)
        else:
            raise ValueError("Unsupported file format. Please upload PDF or DOCX.")
    
    @staticmethod
    def parse_pdf(file):
        """Extract text from PDF file"""
        try:
            # Handle different file types
            if hasattr(file, 'temporary_file_path'):
                # File is stored temporarily on disk
                file_path = file.temporary_file_path()
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + '\n'
                    return text.strip()
            
            elif hasattr(file, 'read'):
                # File is in memory (UploadedFile)
                file.seek(0)  # Go to beginning of file
                reader = PyPDF2.PdfReader(file)
                text = ''
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + '\n'
                return text.strip()
            
            else:
                # Assume it's a file path string
                with open(file, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + '\n'
                    return text.strip()
                    
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file):
        """Extract text from DOCX file"""
        try:
            if hasattr(file, 'read'):
                # File is in memory (UploadedFile)
                file.seek(0)  # Go to beginning of file
                doc = docx.Document(file)
                text = ''
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + '\n'
                return text.strip()
            
            else:
                # Assume it's a file path string
                doc = docx.Document(file)
                text = ''
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + '\n'
                return text.strip()
                
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")